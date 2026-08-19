#!/usr/bin/env python3
"""Build the downloadable CV (.docx and .pdf) from a single content source.

Content is merged from the AIS resume (Aug 2026) into the academic CV that the
site already carried. The home address that appears on the AIS resume is
deliberately omitted: this file is published on a public website.

Usage:  python3 tools/build_cv.py
Writes: assets/Sambold_Daniel_CV.docx
        assets/Sambold_Daniel_CV.pdf
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from facts import SHARK, REID, RELAY, JUE, HABHUB  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "assets"

NAME = "DANIEL SAMBOLD"
CONTACT = [
    "(858) 610-6227  |  daniel.sambold@gmail.com",
    "Portfolio: https://dbold23.github.io  |  GitHub: github.com/dbold23  |  LinkedIn: linkedin.com/in/daniel-sambold-620b37221",
]

# Each section: (HEADING, [entries])
# entry = {"title", "meta" (optional right-aligned date), "sub" (optional), "bullets": [...]}
# A section may instead be {"lines": [(label, text), ...]} for the skills block.

SECTIONS = [
    (
        "EDUCATION",
        [
            {
                "title": "California State University, Monterey Bay",
                "meta": "Graduated Spring 2026",
                "sub": "B.S. Marine Science, Minor in Biology",
                "bullets": [
                    "UROC Innovation Scholar (funded by CSUN HSI Equity Innovation Hub & Apple)",
                ],
            }
        ],
    ),
    (
        "TECHNICAL SKILLS",
        [
            {
                "labeled": [
                    ("Languages", "Python, R, Java 21, SQL/T-SQL, JavaScript/TypeScript, MATLAB, Bash/Unix"),
                    (
                        "ML & Computer Vision",
                        "PyTorch, YOLO detection and YOLOv8-pose keypoints, SAM 2 / SAM 3 / SlimSAM segmentation, "
                        "DINOv2 embeddings, ArcFace metric learning, XGBoost, scikit-learn, CNN-LSTM, "
                        "positive-unlabeled learning, active learning and human-in-the-loop annotation, "
                        "leakage-aware subject-level validation, ONNX Runtime",
                    ),
                    (
                        "Data Engineering",
                        "Microsoft SQL Server, PostgreSQL, SQLite, Flyway migrations, Hibernate ORM/Envers, JPA, JPQL; "
                        "schema design with composite and filtered indexes; Quarkus, JAX-RS/REST, Flask REST APIs; Docker; "
                        "AWS (EC2, S3, SageMaker); distributed work-queue orchestration (lease-based claiming, crash recovery, "
                        "multi-GPU worker pools); Celery/Redis",
                    ),
                    (
                        "Data Formats & Standards",
                        "COCO (compressed-RLE / pycocotools interoperability), YOLO, Pascal VOC, Darwin Core, "
                        "CF-1.8 NetCDF, Movebank, GeoTIFF, GeoJSON, Parquet",
                    ),
                    (
                        "Ocean & Geospatial Data",
                        "NOAA VIIRS / ERDDAP satellite data, NDBC buoy data, GIS spatial analysis, Shapely, pyproj, rasterio; "
                        "acoustic, accelerometer and magnetometer bio-logging tags; VHF telemetry; adaptive signal processing; "
                        "water quality monitoring; BRUV deployment; CFD modeling (OpenFOAM)",
                    ),
                    (
                        "Statistics",
                        "GAMs, Bayesian hierarchical modeling, Gaussian process regression, bootstrap confidence intervals, "
                        "sequential Monte Carlo particle filtering and FFBS smoothing",
                    ),
                    (
                        "Marine Operations",
                        "AAUS Scientific Diver; NAUI Divemaster (143 logged dives); MOTC Certified Vessel Operator",
                    ),
                ]
            }
        ],
    ),
    (
        "RESEARCH & ENGINEERING EXPERIENCE",
        [
            {
                "title": "Summer Intern, Bioinspiration Lab (FathomNet)",
                "meta": "Jun – Aug 2026",
                "sub": "Monterey Bay Aquarium Research Institute (MBARI), Moss Landing, CA",
                "bullets": [
                    "Designed and prototyped a 10-migration SQL Server schema evolution with rollback scripts, restructuring "
                    "annotations into an observation → localization → geometry hierarchy so multiple contributors' geometries "
                    "attach to one animal while license, observer and review state stay per contribution.",
                    "Extended FathomNet's annotation data model and REST API from bounding boxes only to point and COCO "
                    "compressed-RLE segmentation geometry across JPA entities, repositories, services and resources, porting "
                    "pycocotools' RLE codec to Java with 17 unit tests so masks from SAM or COCO tooling round-trip byte for byte "
                    "with no new dependency.",
                    "Extended MBARI's open-source FathomNet Python client to export point and segmentation annotations alongside "
                    "bounding boxes across COCO, YOLO and Pascal VOC, with per-record license propagation and a gate for datasets "
                    "whose terms are unknown.",
                    "Converted existing bounding-box annotations into graded segmentation masks across the FathomNet database, "
                    "scaling from a SAM 2 pilot to a full-corpus SAM 3 run. Built the sharded, crash-safe orchestrator: shard by "
                    "image, a WAL-mode SQLite manifest as a lock-free work queue with BEGIN IMMEDIATE lease claims so a dead "
                    "worker's shards are reclaimed automatically, temp-and-rename commits so a crash never marks a half-written "
                    "shard complete, and one long-lived worker per GPU with the model loaded once.",
                    "Graded every mask rather than accepting it: best-of-three multimask selection by SAM's predicted IoU, a "
                    "stability score computed from mask logits, and AUTO_ACCEPT / REVIEW / REJECT tiers on dual thresholds, with "
                    "rejects routed to a sidecar carrying an explicit deferred count so accept rate could not be inflated.",
                    "Built a loopback ONNX segmentation service (SlimSAM-77-uniform, ONNX Runtime / CoreML) filling FathomNet's "
                    "SegmentationProvider seam so the editor's Smart Select returns real object masks instead of the shipped "
                    "stub's disc.",
                    "Built a dataset planning tool that characterizes a proposed ML training slice of an annotated survey imagery "
                    "corpus from bounded SQL aggregates: an exclusion waterfall attributing every dropped annotation to one named "
                    "stage, per-concept trainability thresholds and the untrainable long tail, train/test leakage counts under "
                    "four candidate split keys, and unlabelled-positive contamination.",
                ],
            },
            {
                "title": "Lead Developer, anchor: Biologging Trajectory Reconstruction & Behavior Classification",
                "meta": "May – Aug 2026",
                "sub": "Jorgensen Lab, CSUMB  |  co-developed with D. Moran",
                "bullets": [
                    "Built the validation framework checking biologger-derived kinematics against independent optical ground "
                    "truth, converting DeepLabCut and SLEAP keypoint exports into tailbeat frequency, amplitude and body lengths "
                    "per second across four species, with drone and aquarium video speed calibration and paired multi-tag runs "
                    "scored by Cohen's kappa.",
                    "Rebuilt behavior-classifier evaluation for subject-level leave-one-out cross-validation with per-class "
                    "precision/recall/F1, balanced accuracy and macro-F1, replacing an evaluation path that leaked "
                    "individual-animal signal between train and test.",
                    "Implemented verified-position correction threading GPS and acoustic fixes through a particle filter and FFBS "
                    "smoother to arrest dead-reckoning drift, reported with held-out drift RMSE; added CF-1.8 NetCDF, Movebank "
                    "and GeoJSON interchange.",
                ],
            },
            {
                "title": "Undergraduate Researcher, Ocean Predator Ecology Lab, CSUMB",
                "meta": "Spring 2023 – Present",
                "sub": "Mentor: Dr. Salvador Jorgensen",
                "bullets": [
                    f"Developed multiple computer vision pipelines for pose detection and automated shark morphometric "
                    f"analysis across {SHARK['videos']} archival videos ({SHARK['corpus_size']}, {SHARK['corpus_years']}, "
                    f"{SHARK['site_count']} sites), reaching {SHARK['box_map50']} {SHARK['box_map50_label']}, with a SQLite "
                    f"backend structured around Darwin Core concepts.",
                    "Built an active learning pipeline with an iterative human-in-the-loop platform for pose, segmentation and "
                    "3D interpretation tasks.",
                    f"Built SharkScarAnnotator, a multi-annotator platform whose {SHARK['annotator_skeleton_points']}-point "
                    f"labelling skeleton feeds the {SHARK['pose_keypoints']}-keypoint pose model, with consensus scoring "
                    f"weighted by annotator experience and confidence.",
                    "Processing accelerometer and magnetometer sensor data alongside dorsal-mounted camera video to correlate "
                    "movement signatures with observed white shark behaviors.",
                    "Performed computational fluid dynamics analysis (OpenFOAM) of acoustic tags mounted to shark dorsal fins, "
                    "iterating tag geometry to minimize hydrodynamic drag.",
                ],
            },
            {
                "title": "Research Assistant, Jue Lab, CSUMB",
                "meta": "Spring 2025 – Present",
                "sub": "Mentor: Dr. Nathaniel Jue",
                "bullets": [
                    f"Built a {JUE['pipeline_stages']}-stage automated analysis pipeline for TECAN plate reader data covering "
                    f"{JUE['strains']} bacterial strains across {JUE['groups']} experimental groups, assessing pesticide "
                    f"bioremediation potential.",
                    "Implemented modified Gompertz growth modeling, Haldane substrate inhibition kinetics, a HistGBT classifier "
                    "for automated curve quality screening, and Bayesian hierarchical modeling with bootstrap confidence "
                    f"intervals; trained the curve-quality classifier on {JUE['train_synthetic']} synthetic and "
                    f"{JUE['train_real']} audited real curves, then validated it against an independent "
                    f"{JUE['validation_curves']}-curve synthetic suite at {JUE['validation_accuracy']} accuracy.",
                ],
            },
            {
                "title": "Lead Technician, Aquaculture Lab, Moss Landing Marine Laboratories",
                "meta": "Spring 2024 – Present",
                "sub": "Mentor: Dr. Luke Gardner",
                "bullets": [
                    "Independently designing and managing two purple sea urchin (Strongylocentrotus purpuratus) aquaculture "
                    "experiments, stocking density optimization and finishing feed roe enhancement, across the full lifecycle "
                    "from hypothesis through manuscript; two manuscripts in review.",
                    "Performing Texture Profile Analysis (TPA) to assess gonad quality and commercial viability; contributing to "
                    "kelp forest restoration through sustainable urchin ranching.",
                ],
            },
        ],
    ),
    (
        "SELECTED PROJECTS",
        [
            {
                "title": "HABHub: Harmful Algal Bloom Forecasting Platform",
                "meta": "1st Place, NOAA SatHack 2025",
                "bullets": [
                    "Real-time HAB monitoring integrating NOAA VIIRS satellite imagery via ERDDAP, NDBC buoy data and CalHABMAP "
                    f"toxin records for Monterey Bay; {HABHUB['features']}-feature XGBoost model behind {HABHUB['apis']}, with "
                    f"{HABHUB['dashboards']}, plus automated alerting; containerized production deploy.",
                ],
            },
            {
                "title": "Individual Re-Identification Across Species",
                "meta": "2025 – Present",
                "bullets": [
                    f"One re-identification approach applied to more than one animal: SAM 2 segmentation, feature "
                    f"embedding through an ArcFace metric head, and cosine-similarity ranking against a gallery of "
                    f"known individuals, so photographic sightings support mark-recapture without physical tags.",
                    f"Harbor porpoise, by dorsal fin outline. {REID['porpoise_backbone']} embeddings across "
                    f"{REID['porpoise_individuals']} individuals ({REID['porpoise_sightings']} sightings, "
                    f"{REID['porpoise_images']} images), after benchmarking {REID['porpoise_backbones_evaluated']}. "
                    f"Unbiased temporal split ({REID['porpoise_temporal_condition']}): "
                    f"{REID['porpoise_temporal_rank1']} rank-1, {REID['porpoise_temporal_rank5']} rank-5. "
                    f"Leave-one-out on {REID['porpoise_loo_condition']} reaches {REID['porpoise_loo_rank1']} rank-1 but is "
                    f"inflated by design and reported only for the gallery-size experiment.",
                    f"Broadnose sevengill shark (<em>{REID['sevengill_species']}</em>), by spot constellation rather than "
                    f"outline: {REID['sevengill_images']} images catalogued, {REID['sevengill_labelled_individuals']} "
                    f"individuals labelled, three tracks in parallel ({REID['sevengill_tracks']}). At this data scale "
                    f"local-feature aggregation outperforms fine-tuned metric learning, so the zero-shot track leads.",
                    f"Catalogue platform built to host {REID['platform_scope']}. Gradio interface, SQLite catalog, "
                    f"Docker deployment.",
                ],
            },
            {
                "title": "VHF Wildlife Tag Detection System",
                "meta": "2024 – Present",
                "bullets": [
                    f"Raspberry Pi 5, RTL-SDR Blog V4 and SIM7028 NB-IoT receiver with adaptive pulse detection, "
                    f"auto-calibration and field-tested gain sweep optimization. Measured {RELAY['max_range']} maximum range at "
                    f"a {RELAY['detection_rate']} detection rate for {RELAY['cost']} in components, against a "
                    f"{RELAY['commercial_receiver']} commercial receiver, adding autonomous logging and cellular relay.",
                ],
            },
            {
                "title": "Autonomous Aerial Platform for Marine Survey",
                "meta": "2024 – Present",
                "bullets": [
                    "Custom FPV aircraft for testing computer vision detection from aerial imagery; flight controller "
                    "programming, motor/ESC configuration and radio telemetry systems.",
                ],
            },
        ],
    ),
    (
        "PUBLICATIONS",
        [
            {
                "bullets": [
                    "Sambold, D., & Gardner, L. Effects of Stocking Density on Gonad Development and Survival in Purple Sea "
                    "Urchin (Strongylocentrotus purpuratus) Aquaculture. In review.",
                    "Sambold, D., & Gardner, L. Examining Finishing Feed Applications for Purple Urchin "
                    "(Strongylocentrotus purpuratus) Roe Enhancement. In review.",
                ]
            }
        ],
    ),
    (
        "PRESENTATIONS",
        [
            {
                "bullets": [
                    "Sambold, D., Gupta, I., & Jorgensen, S. “Assessing Shark Body Condition.” Poster, NEPSS, "
                    "March 2026, Newport, OR.",
                    "Sambold, D., & Gardner, L. “Optimizing Urchin Aquaculture: Carotenoid Application & Pellet "
                    "Processing.” Aquaculture America, February 2026, Las Vegas, NV.",
                    "Sambold, D., Gupta, I., & Jorgensen, S. “Assessing White Shark Body Condition.” Poster, Western "
                    "Society of Naturalists, November 2025, San Diego, CA.",
                    "Sambold, D., & Jorgensen, S. “Classifying White Shark Behaviors from Sensor Packages.” UROC Summer "
                    "Symposium, CSUMB, Summer 2025.",
                    "Sambold, D. “From Fins to Frames: An AI-Powered Tool for Morphometric Shark Analysis.” UROC Spring "
                    "Showcase, CSUMB, Spring 2025.",
                    "Sambold, D. “HAB Hub: Integrating Satellite, Buoy, and Toxin Data for Real-Time HAB Monitoring in "
                    f"Monterey Bay.” ESIP, {HABHUB['esip_date']}, Online.",
                ]
            }
        ],
    ),
    (
        "FIELD EXPERIENCE & CERTIFICATIONS",
        [
            {
                "title": "Research & Diver Development Program, South Africa",
                "meta": "July – August 2024",
                "bullets": [
                    "31 hours of scientific diving: nudibranch population surveys, swell shark egg casing measurement, BRUV "
                    "deployments, fish dissections and quadrat-based invertebrate counts, microplastic sampling.",
                    "Advanced vessel operations: docking, loading, anchoring, navigation, and equipment deployment and recovery "
                    "at sea.",
                ],
            },
            {
                "title": "AAUS Scientific Diver, CSUMB",
                "meta": "Spring 2024",
                "bullets": [
                    "Designed and conducted “Thermal Thresholds: The Role of Temperature Fluctuations in Shaping Kelp Forest "
                    "Health at San Carlos Beach, Monterey, CA.”",
                ],
            },
            {
                "title": "NAUI Divemaster, 143 logged dives",
                "meta": "July 2024 – Present",
                "bullets": [
                    "Certified professional dive leader; dive planning, safety management and underwater research coordination.",
                ],
            },
            {
                "title": "MOTC Certified Vessel Operator, CSUMB",
                "meta": "2024",
                "bullets": [],
            },
            {
                "title": "CITI Certified, Responsible Conduct of Research",
                "meta": "",
                "bullets": [],
            },
        ],
    ),
    (
        "LEADERSHIP & OUTREACH",
        [
            {
                "title": "Trail Committee Chair, Advocates for Nisene Marks State Park",
                "meta": "2024 – 2025",
                "bullets": [
                    "Led hazard assessments with 30+ volunteers; developed GIS-based stewardship mapping for park management.",
                ],
            },
            {
                "title": "Educational Aide, Elkhorn Slough Foundation",
                "meta": "2024 – 2025",
                "bullets": [
                    "Led environmental education programs for students in grades 4–10; habitat restoration and community "
                    "science.",
                ],
            },
            {
                "title": "Relay Station Training Program, Galápagos & California",
                "meta": "2024 – Present",
                "bullets": [
                    f"Designed and built a low-cost {RELAY['band']} telemetry relay station, deployed it in "
                    f"{RELAY['sites']}, and taught local teams to build, site and repair their own.",
                ],
            },
        ],
    ),
]



def plain(text):
    """Inline HTML is for the web renderer only; docx and pdf take plain text."""
    import re
    return re.sub(r'</?(?:em|strong)>', '', text)

# ---------------------------------------------------------------- DOCX ------

def build_docx(path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ACCENT = RGBColor(0x0A, 0x36, 0x63)
    RIGHT_TAB = Inches(7.0)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")

    def para(space_before=0, space_after=0, tab_right=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if tab_right:
            p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        return p

    def bottom_border(p):
        ppr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0A2463")
        borders.append(bottom)
        ppr.append(borders)

    # Header
    p = para(space_after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT

    for line in CONTACT:
        p = para(space_after=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(8.5)

    for heading, entries in SECTIONS:
        p = para(space_before=9, space_after=3)
        r = p.add_run(heading)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = ACCENT
        bottom_border(p)

        for entry in entries:
            if "labeled" in entry:
                for label, text in entry["labeled"]:
                    p = para(space_after=2)
                    p.paragraph_format.left_indent = Inches(0.14)
                    p.paragraph_format.first_line_indent = Inches(-0.14)
                    rl = p.add_run(f"{label}: ")
                    rl.bold = True
                    p.add_run(plain(text))
                continue

            if entry.get("title"):
                p = para(space_before=4, space_after=0, tab_right=bool(entry.get("meta")))
                r = p.add_run(entry["title"])
                r.bold = True
                r.font.size = Pt(10)
                if entry.get("meta"):
                    p.add_run("\t")
                    rm = p.add_run(entry["meta"])
                    rm.italic = True
                    rm.font.size = Pt(9)

            if entry.get("sub"):
                p = para(space_after=1)
                r = p.add_run(entry["sub"])
                r.italic = True
                r.font.size = Pt(9)

            for b in entry.get("bullets", []):
                p = para(space_after=1)
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.first_line_indent = Inches(-0.14)
                p.add_run("•  " + plain(b))

    doc.save(str(path))
    return path


# ----------------------------------------------------------------- PDF ------

def build_pdf(path):
    from fpdf import FPDF

    FONT_DIR = Path("/System/Library/Fonts/Supplemental")
    ACCENT = (10, 36, 99)

    pdf = FPDF(format="letter", unit="pt")
    pdf.set_margins(54, 36, 54)
    pdf.set_auto_page_break(True, margin=36)
    pdf.add_font("Body", "", str(FONT_DIR / "Arial.ttf"))
    pdf.add_font("Body", "B", str(FONT_DIR / "Arial Bold.ttf"))
    pdf.add_font("Body", "I", str(FONT_DIR / "Arial Italic.ttf"))
    pdf.add_page()

    width = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.set_font("Body", "B", 19)
    pdf.set_text_color(*ACCENT)
    pdf.cell(width, 24, NAME, align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Body", "", 7.6)
    pdf.set_text_color(60, 60, 60)
    for line in CONTACT:
        pdf.cell(width, 10, line, align="C", new_x="LMARGIN", new_y="NEXT")

    BULLET_H = 9.3
    BULLET_PT = 7.7

    def bullet_height(text, indent):
        """Rendered height of one bullet, so headers can be kept with their body."""
        pdf.set_font("Body", "", BULLET_PT)
        lines = pdf.multi_cell(
            width - indent, BULLET_H, text, align="L", dry_run=True, output="LINES"
        )
        return len(lines) * BULLET_H

    for heading, entries in SECTIONS:
        pdf.ln(6)
        pdf.set_font("Body", "B", 9.5)
        pdf.set_text_color(*ACCENT)
        # Never strand a section heading at the foot of a page
        if pdf.will_page_break(34):
            pdf.add_page()
        pdf.cell(width, 12, heading, new_x="LMARGIN", new_y="NEXT")
        y = pdf.get_y()
        pdf.set_draw_color(*ACCENT)
        pdf.set_line_width(0.6)
        pdf.line(pdf.l_margin, y, pdf.l_margin + width, y)
        pdf.ln(3)
        pdf.set_text_color(20, 20, 20)

        for entry in entries:
            if "labeled" in entry:
                for label, text in entry["labeled"]:
                    pdf.set_font("Body", "B", BULLET_PT)
                    label_w = pdf.get_string_width(label + ": ")
                    start_y = pdf.get_y()
                    pdf.cell(label_w, 9.4, label + ": ")
                    pdf.set_font("Body", "", BULLET_PT)
                    pdf.set_xy(pdf.l_margin + label_w, start_y)
                    pdf.multi_cell(width - label_w, 9.4, plain(text), align="L",
                                   new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(1)
                continue

            if entry.get("title"):
                meta = entry.get("meta") or ""
                pdf.set_font("Body", "I", 8)
                meta_w = pdf.get_string_width(meta) + 2 if meta else 0
                pdf.set_font("Body", "B", 8.8)
                title_lines = pdf.multi_cell(
                    width - meta_w, 11, entry["title"], align="L",
                    dry_run=True, output="LINES",
                )
                # Keep the header with its subtitle and first bullet, otherwise the
                # right-aligned date lands on a different page than its title
                keep = len(title_lines) * 11
                if entry.get("sub"):
                    keep += 10
                if entry.get("bullets"):
                    keep += bullet_height(entry["bullets"][0], 15)
                if pdf.will_page_break(keep):
                    pdf.add_page()

                pdf.set_font("Body", "B", 8.8)
                start_y = pdf.get_y()
                pdf.multi_cell(width - meta_w, 11, entry["title"], align="L",
                               new_x="LMARGIN", new_y="NEXT")
                if meta:
                    end_y = pdf.get_y()
                    pdf.set_font("Body", "I", 8)
                    pdf.set_text_color(90, 90, 90)
                    pdf.set_xy(pdf.l_margin + width - meta_w, start_y)
                    pdf.cell(meta_w, 11, meta, align="R")
                    pdf.set_text_color(20, 20, 20)
                    pdf.set_xy(pdf.l_margin, end_y)

            if entry.get("sub"):
                pdf.set_font("Body", "I", 7.7)
                pdf.set_text_color(70, 70, 70)
                pdf.multi_cell(width, 10, entry["sub"], align="L",
                               new_x="LMARGIN", new_y="NEXT")
                pdf.set_text_color(20, 20, 20)

            for b in entry.get("bullets", []):
                pdf.set_font("Body", "", BULLET_PT)
                start_y = pdf.get_y()
                pdf.set_xy(pdf.l_margin + 6, start_y)
                pdf.cell(9, BULLET_H, "•")
                pdf.set_xy(pdf.l_margin + 15, start_y)
                pdf.multi_cell(width - 15, BULLET_H, plain(b), align="L",
                               new_x="LMARGIN", new_y="NEXT")
                pdf.ln(0.5)
            pdf.ln(1.5)

    pdf.output(str(path))
    return path




# ----------------------------------------------------------------- HTML -----

def build_html(index_path):
    """Rewrite the CV block inside index.html from the same SECTIONS data.

    The page used to carry a hand-typed copy of the CV, which is how the site
    ended up stating the same metric three different ways. Now there is one
    source and three renderers.
    """
    import html as _html
    import re

    L = []
    a = L.append
    ind = ' ' * 12

    a(f'{ind}<h2>Resume / CV</h2>')
    a(f'{ind}<p class="cv-contact">'
      f'<a href="mailto:daniel.sambold@gmail.com">daniel.sambold@gmail.com</a>'
      f' &middot; <a href="https://github.com/dbold23" target="_blank" rel="noopener">github.com/dbold23</a>'
      f' &middot; <a href="https://www.linkedin.com/in/daniel-sambold-620b37221" target="_blank" rel="noopener">LinkedIn</a>'
      f'</p>')
    a(f'{ind}<p class="cv-downloads">')
    a(f'{ind}    <a href="assets/Sambold_Daniel_CV.pdf" download class="cv-download-btn">Download CV (PDF)</a>')
    a(f'{ind}    <a href="assets/Sambold_Daniel_CV.docx" download class="cv-download-alt">or .docx</a>')
    a(f'{ind}</p>')

    def esc(t):
        # the source strings are plain text; keep the few tags we do want
        t = _html.escape(t, quote=False)
        t = t.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
        t = t.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
        return t

    for heading, entries in SECTIONS:
        a(f'{ind}<h3>{esc(heading.title())}</h3>')
        for entry in entries:
            a(f'{ind}<div class="resume-item">')
            if 'labeled' in entry:
                a(f'{ind}    <ul>')
                for label, text in entry['labeled']:
                    a(f'{ind}        <li><strong>{esc(label)}:</strong> {esc(text)}</li>')
                a(f'{ind}    </ul>')
            else:
                if entry.get('title'):
                    a(f'{ind}    <h4>{esc(entry["title"])}</h4>')
                meta_bits = [b for b in (entry.get('meta'), entry.get('sub')) if b]
                if meta_bits:
                    a(f'{ind}    <p class="resume-meta">{esc(" | ".join(meta_bits))}</p>')
                if entry.get('bullets'):
                    a(f'{ind}    <ul>')
                    for b in entry['bullets']:
                        a(f'{ind}        <li>{esc(b)}</li>')
                    a(f'{ind}    </ul>')
            a(f'{ind}</div>')

    block = '\n'.join(L)
    text = index_path.read_text()
    start = '<!-- CV:START generated by tools/build_cv.py, do not edit by hand -->'
    end = '<!-- CV:END -->'
    pattern = re.compile(re.escape(start) + '.*?' + re.escape(end), re.S)
    if not pattern.search(text):
        raise SystemExit('CV markers not found in index.html')
    index_path.write_text(pattern.sub(f'{start}\n{block}\n{ind}{end}', text))
    return len(L)


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True)
    d = build_docx(OUT / "Sambold_Daniel_CV.docx")
    p = build_pdf(OUT / "Sambold_Daniel_CV.pdf")
    for f in (d, p):
        print(f"{f.name}: {f.stat().st_size:,} bytes")
    n = build_html(ROOT / "index.html")
    print(f"index.html CV block: {n} lines regenerated")
